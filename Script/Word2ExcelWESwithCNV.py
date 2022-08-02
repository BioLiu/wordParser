#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Tue Oct 19 22:34:23 2021

@author: chenliu

"""

import os
import argparse
import re
from docx import Document

# Function Objects2list takes contents from objects and puts them into a list
def Objects2list(objects):
    li = []
    for i in objects:
        if i.text.strip() == "" or i.text.strip().startswith("家系验证") or i.text.strip().startswith("检出"):
            next
        else:
            li.append(i.text.strip())
    return(li)

def Reference2list(objects):
    li = []
    for i in objects:
        if i.text.strip() == "":
            next
        else:
            li.append(i.text.strip())
    return(li)

def ObjectsBreak2list(objects):
    pgh_li = []
    for i in objects:
        if i.text == "" or i.text.strip().startswith("家系验证") or i.text.strip().startswith("检出"):
            next
        else:
            pgh_li.append(i)
    
    li = []    
    for i in range(len(pgh_li)):
        text = pgh_li[i].text
        if text.startswith("  "):
            if text.strip().endswith("。"):
                li.append(text.strip())
        elif pgh_li[i].paragraph_format.first_line_indent:
            if pgh_li[i].paragraph_format.first_line_indent > 0:
                if text.strip().endswith("。"):
                    li.append(text.strip())
        else:
            textmerge = pgh_li[i-1].text.strip()+text.strip()
            li.append(textmerge)        
    return(li)

def get_filelist(inputdir): 
    Filelist = [] 
    for home, dirs, files in os.walk(inputdir): 
        for filename in files: 
            # 文件名列表，包含完整路径
            if "docx" in filename:
                Filelist.append(os.path.join(home, filename))
                # # 文件名列表，只包含文件名     
                # Filelist.append( filename)
            else:
                print(filename) 
    return Filelist


def separatorParser(text):
    '''
    Identify the separator used in the Word ducument and split contents with that separator.
    '''
    
    textout = "-"
    if ":" in text:
        textout = text.split(":")[1].strip()
    elif "：" in text:
        textout = text.split("：")[1].strip()
    return textout

def WordParser(inputdir, outputfile):
    classmarker = "=============" + outputfile.split("/")[-1][:-4] + "============="
    print(classmarker)
    Filelist   = get_filelist(inputdir)
    for inputfile in Filelist:
        start_info = "Processing on %s" %(inputfile.split("/")[-1].split(".")[0])
        print(start_info)
        fo = open(outputfile,"a")
        try:
            document = Document(inputfile)    
            # Pick up information from tables in the word file        
            table_list = document.tables
            proband_tb = table_list[0]
            name       = separatorParser(proband_tb.rows[0].cells[0].text)
            patient_no = separatorParser(proband_tb.rows[0].cells[2].text)
            recall_no  = separatorParser(proband_tb.rows[0].cells[-1].text.split("  ")[-1])
            sex        = separatorParser(proband_tb.rows[1].cells[0].text)
            department = separatorParser(proband_tb.rows[1].cells[2].text)
            diagnosis  = separatorParser(proband_tb.rows[1].cells[3].text)
            age        = separatorParser(proband_tb.rows[2].cells[0].text)
            sample_no  = separatorParser(proband_tb.rows[2].cells[1].text)
            doctor     = separatorParser(proband_tb.rows[2].cells[2].text)
            birthday   = "-"
            variant_info = "\t".join((['-']*8))
            diseases_omim = "-"
            diseases_pattern = "-"
            variant_conclusion = "-"
            cnv_conclusion = "-"
            variant_explain = "-"
            variant_references = "-"
            ClinVar    = "-"
            HGMD       = "-"
            LOVD       = "-"
            
            
            # Patients with mutations
            mutation_tb = False
            cnv_tag = False
            for i in range(len(table_list)):
                text = table_list[i].rows[0].cells[0].text.strip()
                if text == "基因":
                    mutation_tb = table_list[i]
                elif "样本采集时间" in text:
                    report_tb = table_list[i]
            
            receive_time  = separatorParser(report_tb.rows[0].cells[1].text)
            
            # Analyze paragraphs
            paragraph_list     = document.paragraphs
            for i in range(len(paragraph_list)):
                content = document.paragraphs[i].text
                if "临床信息：" in content:
                    clinical_info = document.paragraphs[i+1].text.strip()
                    pbd = re.compile(r'出生日期(.{4,15})[）]')
                    if pbd.search(clinical_info):
                        birthday = pbd.search(clinical_info).group(1)
                
            
            # 提取拷贝数变异信息
            paragraph_list     = document.paragraphs
            for i in range(len(paragraph_list)):
                content = document.paragraphs[i].text
                if "（主要分析CNV）" in content or "检测结果:" in content:
                    cnv_start = i+1
                if "检测结论：" in content:
                    cnv_end = i
                if "深度数据挖掘分析，" in content:
                    cnv_tag = True
                    # print(content)
                    cnv_patterns = re.compile(r"显示[\u4e00-\u9fa5]{2,10}(.*)的[\u4e00-\u9fa5]*捕获测序")
                    fg_match = cnv_patterns.search(content)                    
                    # if fg_match:
                    #     cnv_conclusion = content.strip()
                    #     # fg_pos_patterns = re.compile(r"")
            
            if mutation_tb:
                reference = False
                comment_status = False
                for i in range(len(paragraph_list)):
                    content = document.paragraphs[i].text
                    if "检测结论：" in content:
                        conclusion_start = i+1
                    if "结果解释：" in content or "检测解释：" in content:
                        conclusion_end = i
                        explain_start  = i+1
                    if "咨询建议：" in content:
                        explain_end = i
                    if "参考资料：" in content or "参考文献：" in content or "参考资料及文献：" in content:
                        reference_start = i+1
                        reference = True
                    if "备注：" in content:
                        reference_end = i
                        comment_status = True
                    if not comment_status:
                        if "附一：" in content:
                            reference_end = i
                        
                            
                conclusion_li = Objects2list(paragraph_list[conclusion_start:conclusion_end])
                explain_li    = ObjectsBreak2list(paragraph_list[explain_start:explain_end])
                if reference:
                    reference_li  = Reference2list(paragraph_list[reference_start:reference_end])
                    num_li = []
                    for num in range(15):
                        num_li.append(str(num))
                    reference_dic = {}
                    for i in range(len(reference_li)):
                        text = reference_li[i].strip()
                        if text.startswith("["):
                            reference_dic[str(i+1)] = text
                        else: 
                            if len(reference_li) < 10:
                                if text[0] in num_li:
                                    reference_dic[text[0]] = text
                                else:
                                    reference_dic[str(i+1)]  = "[%s]"%(str(i+1)) + " "+ text
                            else:
                                reference_dic[str(i+1)]  = "[%s]"%(str(i+1)) + " "+ text
                # Mutation table
                mutation_tb      = table_list[1]
                mutation_objects = mutation_tb.columns[0].cells[1:]
                mutation_li      = Objects2list(mutation_objects)
                mutation_no      = len(mutation_li)
                mutation_no_uniq = len(list(set(mutation_li)))
                column_no    = len (mutation_tb.rows[0].cells)
                for i in range(len(mutation_li)):
                    identity = i+1
                    variant_objects  =  mutation_tb.rows[identity].cells
                    variant_gene     = variant_objects[0].text.strip()
                    variant_gene     = variant_gene.replace("\n","")
                    variant_gene     = variant_gene.replace("\r","")
                    variant_dna      = variant_objects[4].text.strip()
                    variant_dna      = variant_dna.replace("\r","")
                    variant_dna      = variant_dna.replace("\n","")
                    variant_dna      = variant_dna.replace(" ","")
                    if "/" in variant_dna:
                        variant_dna = variant_dna.split("/")[0]
                    variant_info     = "\t".join((['-']*9))
                    diseases_omim    = "-"
                    diseases_pattern = "-"
                    variant_conclusion = "-"
                    variant_explain = "-"
                    gene_omim = "-"
                    variant_references = "-"
                    ClinVar    = "-"
                    HGMD       = "-"
                    LOVD       = "-"
                    
                    if "(" in variant_dna:
                        variant_dna = variant_dna.split("(")[0]
                    
                    for conclusion in conclusion_li:
                        conclusion = conclusion.replace(" ","")
                        if variant_dna in conclusion:
                            variant_conclusion = conclusion
                    
                    for explain in explain_li:
                        explain = explain.replace(" ","")
                        if explain.startswith("OMIM") and variant_gene in explain:
                            gene_omim = explain
                        elif explain.startswith(variant_gene) and variant_dna in explain:
                            variant_explain = gene_omim + "|" + explain
                    
                    # pdb: pattern for the database
                    pdb = re.compile(r"[a-z,A-Z]{4,7}.{2,10}注释为“.{2,18}”")
                    po  = re.compile(r"([a-z,A-Z]{4,7}).{2,10}注释为(“.{2,18}”)")
                    database_match_li  =  pdb.findall(variant_explain)
                    if "数据库均有收录" in variant_explain:
            #            if ">" in variant_dna:
            #                variant_dna_org  = variant_dna.split(">")[0]
            #                variant_dna_alt  = variant_dna.split(">")[1]
            #                LOVD = "https://databases.lovd.nl/shared/variants/in_gene?search_geneid=%3D%22%s%22&search_VariantOnTranscript/DNA=%3D%22%s%3E%s%22" %(variant_gene, variant_dna_org, variant_dna_alt)
            #            else:
            #                LOVD = "有"
                        LOVD =     "有"
                        HGMD     = "有"
                        ClinVar  = "有"
                        
                        
                    database_dic = {}
                    for exp in database_match_li:
                        mo = po.search(exp)
                        if mo:
                            database_dic[mo.group(1)] = mo.group(2)[1:-1]
                    if "ClinVar" in database_dic.keys():
                        ClinVar = database_dic["ClinVar"]
                    if "HGMD"    in database_dic.keys():
                        HGMD    = database_dic["HGMD"]
                    if "LOVD"    in database_dic.keys():
                        LOVD    = database_dic["LOVD"]
                    # pr: pattern for the references
                    pr = re.compile(r"\[\d{1}[-,]?\d?]")
                    reference_match = pr.findall(variant_explain)
                    reference_ids    = []
                    for i in reference_match:
                        i = i[1:-1]
                        if "-" in i:
                            end   = int(i.split("-")[-1]) + 1
                            start = int(i.split("-")[0])
                            for j in range(start, end):
                                reference_ids.append(j)
                        elif "," in i:
                            end   = int(i.split(",")[-1]) + 1
                            start = int(i.split(",")[0])
                            for j in range(start, end):
                                reference_ids.append(j)
                        else:
                            reference_ids.append(int(i))
                    reference_ids = list(set(reference_ids))
                    # sort the list
                    reference_ids.sort()
                    # references of the variant
                    if reference:
                        variant_reference_li = []
                        for i in reference_ids:
                            i = str(i)
                            variant_reference_li.append(reference_dic[i])
                        
                        variant_references = "|".join((variant_reference_li))
                    else:
                        variant_references = "-"
                    
                    if column_no <= 10:
                        variant_info =  "\t".join((Objects2list(variant_objects)))
                        # pd: pattern for diseases 
                        pd = re.compile(r"(基因与)(.+)([相|有]关)")
                        # ph: pattern for the heredity
                        ph = re.compile(r"(.{1}染色体.{1,20}遗传)")
                        phlink = re.compile(r"(.{1}连锁.{1,20}遗传)")
                        if pd.search(gene_omim):
                            diseases_omim    = pd.search(gene_omim).group(2)
                            if "性遗传" in gene_omim:
                                if ph.search(gene_omim):
                                    diseases_pattern = ph.search(gene_omim).group(1)
                                else:
                                    if phlink.search(gene_omim):
                                        diseases_pattern = phlink.search(gene_omim).group(1)
                                    else:
                                        diseases_pattern = ""
                                if "。" in diseases_pattern:
                                    diseases_pattern = diseases_pattern.split("。")[0]
                            else:
                                diseases_pattern = ""                    
                  
                    elif column_no > 10:
                        variant_li   = Objects2list(variant_objects)
                        variant_li_temp = variant_li
                        diseases_omim    = variant_li[7].strip()
                        diseases_pattern = variant_li[8].strip()
                        variant_li_temp.pop(7)
                        variant_li_temp.pop(7)
                        variant_info =  "\t".join((variant_li_temp))
                        #print(diseases_omim + "\n")
                    variant_info = variant_info.replace(" ","")
                    if variant_dna != "-" and (variant_conclusion == "-" or variant_explain == "-"):
                        error_info = "Error in the sample %s" %(inputfile.split("/")[-1].split(".")[0])
                        print(error_info + "\n")
                    
                    variant_info_li = [sample_no, name, patient_no, recall_no, sex, age,
                                       birthday, department, doctor, receive_time, 
                                       clinical_info, diagnosis, variant_info, 
                                       diseases_omim, diseases_pattern, variant_conclusion, 
                                       variant_explain, variant_references, 
                                       ClinVar, HGMD, LOVD]
                    variant_info_out = "\t".join((variant_info_li))
                    variant_info_out = variant_info_out.replace("\n","")
                    variant_info_out = variant_info_out.replace("\r","")
                    #fo.write(variant_info_out + "\n")
            if cnv_tag:
                cnv_li = Objects2list(paragraph_list[cnv_start:cnv_end])
                cnv_conclusion_li = []
                for i in cnv_li:
                    cnv_conclusion_li.append(i.strip())
                cnv_conclusion = "|".join((cnv_conclusion_li))
                diseases_omim = "-"
                diseases_pattern = "-"
                variant_explain = "-"
                variant_references = "-"
                ClinVar    = "-"
                HGMD       = "-"
                LOVD       = "-"
                variant_conclusion = cnv_conclusion
                variant_info_li = [sample_no, name, patient_no, recall_no, sex, age,
                   birthday, department, doctor, receive_time, 
                   clinical_info, diagnosis, variant_info, 
                   diseases_omim, diseases_pattern, variant_conclusion, 
                   variant_explain, variant_references, 
                   ClinVar, HGMD, LOVD]
                variant_info_out = "\t".join((variant_info_li))
                variant_info_out = variant_info_out.replace("\n","")
                variant_info_out = variant_info_out.replace("\r","")
                fo.write(variant_info_out + "\n")
                
            else :
                variant_info_li = [sample_no, name, patient_no, recall_no, sex, age,
                   birthday, department, doctor, receive_time, 
                   clinical_info, diagnosis, variant_info, 
                   diseases_omim, diseases_pattern, variant_conclusion, 
                   variant_explain, variant_references, 
                   ClinVar, HGMD, LOVD]
                variant_info_out = "\t".join((variant_info_li))
                variant_info_out = variant_info_out.replace("\n","")
                variant_info_out = variant_info_out.replace("\r","")
                fo.write(variant_info_out + "\n")
            
            fo.close()    
            end_info = "Finish %s" %(inputfile.split("/")[-1].split(".")[0])
            print(end_info)
        except Exception as e:
            print(e)
            pass
            error_smp  = "Invalid sample: %s" %(inputfile.split("/")[-1].split(".")[0])
            print(error_smp + "\n")
    
    
    
parser = argparse.ArgumentParser(description="Convert Word documents into the table.")
parser.add_argument("-d", help="The path of the folder containing Word ducuments.", type=str)
parser.add_argument("-o", help="The path of the output (table file).", type=str)
args = parser.parse_args()

if args.d and args.o:
    inputdir = args.d
    outputfile = args.o   
    WordParser(inputdir, outputfile)
else:
    print("Please provide paths of the folder and the output.")    
    







